from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import docx
import fitz  # PyMuPDF
import io
import os
import logging
from typing import Optional, Dict, Any
import PyPDF2
import sys
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin
import time
import re
import html
from requests.exceptions import RequestException, Timeout
import json
import hashlib
from datetime import datetime
import bleach
import validators
import mimetypes
import youtube_dl
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter

logger = logging.getLogger(__name__)

class DocumentParser:
    # 敏感词列表（示例）
    SENSITIVE_WORDS = {
        'spam', 'scam', 'hack', 'crack',
        # 添加其他敏感词
    }
    
    # 广告相关标识
    AD_PATTERNS = [
        r'advertisement',
        r'sponsored',
        r'promotion',
        r'ads-container',
        r'banner-ads',
    ]

    def __init__(self):
        self._check_dependencies()
        self.session = self._create_session()
        
    def _check_dependencies(self):
        """检查必要的依赖是否已安装"""
        try:
            # 检查 tesseract
            if sys.platform.startswith('win'):
                if not os.environ.get('TESSERACT_CMD'):
                    tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                    if os.path.exists(tesseract_path):
                        pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    else:
                        logger.warning("Tesseract not found in default location")
            
            # 测试 tesseract 是否可用
            try:
                pytesseract.get_tesseract_version()
                logger.info("Tesseract is available")
            except Exception as e:
                logger.error(f"Tesseract is not properly installed: {str(e)}")
            
            # 检查中文语言包
            if sys.platform.startswith('win'):
                tessdata_path = os.path.join(os.path.dirname(pytesseract.pytesseract.tesseract_cmd), 'tessdata')
            else:
                tessdata_path = '/usr/share/tesseract-ocr/4.00/tessdata'
            
            if not os.path.exists(os.path.join(tessdata_path, 'chi_sim.traineddata')):
                logger.warning("Chinese language pack not found")
            
            # 检查 poppler
            try:
                from pdf2image.exceptions import PDFPageCountError
                logger.info("pdf2image (poppler) is available")
            except ImportError as e:
                logger.error(f"pdf2image or poppler is not properly installed: {str(e)}")
            
            # 检查 PyCryptodome
            try:
                from Crypto.Cipher import AES
                logger.info("PyCryptodome is available")
            except ImportError:
                logger.error("PyCryptodome is not properly installed")
                
        except Exception as e:
            logger.error(f"Error checking dependencies: {str(e)}")

    def parse_file(self, file_path: str) -> Optional[str]:
        """解析文件内容"""
        try:
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._parse_docx(file_path)
            elif file_extension == 'txt':
                return self._parse_txt(file_path)
            elif file_extension in ['png', 'jpg', 'jpeg', 'bmp', 'tiff']:
                return self.parse_image(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}", exc_info=True)
            return None

    def _parse_pdf(self, file_path: str) -> Optional[str]:
        """解析PDF文件"""
        try:
            # 首先尝试使用 PyMuPDF
            doc = fitz.open(file_path)
            text_content = []
            
            # 获取页数
            num_pages = len(doc)
            logger.info(f"PDF has {num_pages} pages")
            
            # 提取所有页面的文本
            for page_num in range(num_pages):
                try:
                    page = doc[page_num]
                    text = page.get_text("text")
                    if text.strip():
                        text_content.append(text.strip())
                        logger.debug(f"Extracted text from page {page_num + 1}")
                    else:
                        logger.warning(f"No text found in page {page_num + 1}")
                except Exception as e:
                    logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue

            doc.close()
            
            if text_content:
                full_text = '\n\n'.join(text_content)  # 使用双换行分隔页面
                logger.info(f"Successfully extracted {len(full_text)} characters using PyMuPDF")
                return full_text
            else:
                logger.info("No text extracted using PyMuPDF, trying OCR")
                return self._ocr_pdf(file_path)

        except Exception as e:
            logger.error(f"Error parsing PDF with PyMuPDF: {str(e)}", exc_info=True)
            logger.info("PyMuPDF failed, trying OCR")
            return self._ocr_pdf(file_path)

    def _parse_docx(self, file_path: str) -> Optional[str]:
        """解析DOCX文件"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # 提取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # 提取所有表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}", exc_info=True)
            return None

    def _parse_txt(self, file_path: str) -> Optional[str]:
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                logger.info(f"Successfully read {len(content)} characters from TXT")
                return content
        except UnicodeDecodeError:
            try:
                # 如果UTF-8解码失败，尝试使用GBK编码
                with open(file_path, 'r', encoding='gbk') as file:
                    content = file.read()
                    logger.info(f"Successfully read {len(content)} characters from TXT using GBK encoding")
                    return content
            except Exception as e:
                logger.error(f"Error parsing TXT file with GBK encoding: {str(e)}", exc_info=True)
                return None
        except Exception as e:
            logger.error(f"Error parsing TXT file: {str(e)}", exc_info=True)
            return None

    def parse_image(self, file_path: str) -> str:
        """解析图片文件"""
        try:
            image = Image.open(file_path)
            # 使用中文和英文语言包
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            return text
        except Exception as e:
            logger.error(f"Error parsing image: {str(e)}", exc_info=True)
            raise 

    def _ocr_pdf(self, file_path: str) -> Optional[str]:
        """使用OCR处理PDF文件"""
        try:
            # 将PDF转换为图片
            logger.info("Converting PDF to images...")
            images = convert_from_path(
                file_path,
                dpi=300,  # 提高DPI以获得更好的文本识别效果
                fmt='jpeg',
                thread_count=os.cpu_count() or 1  # 使用多线程加速
            )
            logger.info(f"Converted {len(images)} pages to images")
            
            # 对每个页面进行OCR
            text_content = []
            for i, image in enumerate(images):
                try:
                    # 预处理图像以提高OCR质量
                    image = Image.fromarray(self._preprocess_image(image))
                    
                    # 使用中文和英文语言包
                    text = pytesseract.image_to_string(
                        image,
                        lang='chi_sim+eng',
                        config='--psm 1 --oem 1'  # 使用更准确的OCR模式
                    )
                    
                    if text.strip():
                        text_content.append(text.strip())
                        logger.debug(f"OCR extracted text from page {i + 1}")
                    else:
                        logger.warning(f"OCR found no text in page {i + 1}")
                except Exception as e:
                    logger.error(f"Error performing OCR on page {i + 1}: {str(e)}")
                    continue
            
            if not text_content:
                logger.error("No text could be extracted from PDF using either method")
                return None
            
            full_text = '\n\n'.join(text_content)  # 使用双换行分隔页面
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF using OCR")
            return full_text
        except Exception as e:
            logger.error(f"Error performing OCR on PDF: {str(e)}", exc_info=True)
            return None

    def _preprocess_image(self, image):
        """预处理图像以提高OCR质量"""
        import numpy as np
        from PIL import ImageEnhance
        
        # 转换为numpy数组
        img_array = np.array(image)
        
        # 转换为PIL图像进行增强
        img = Image.fromarray(img_array)
        
        # 增加对比度
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # 增加锐度
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.5)
        
        return np.array(img) 

    def _create_session(self) -> requests.Session:
        """创建具有重试机制的会话"""
        session = requests.Session()
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)
        return session

    def parse_url(self, url: str) -> Optional[str]:
        """从 URL 解析内容"""
        try:
            logger.info(f"Starting to parse URL: {url}")
            
            # URL验证
            if not self._validate_url(url):
                raise ValueError("Invalid URL format or potentially dangerous URL")
            
            # 获取响应
            response = self._get_response(url)
            
            # 解析内容
            content = self._parse_web_content(response)
            
            # 内容过滤和清理
            content = self._filter_content(content)
            
            # 返回处理后的文本
            return content
            
        except Exception as e:
            self._handle_error(e)
            return None

    def _validate_url(self, url: str) -> bool:
        """验证URL"""
        if not validators.url(url):
            return False
            
        parsed_url = urlparse(url)
        
        # 检查协议
        if parsed_url.scheme not in ['http', 'https']:
            return False
            
        # 检查危险域名（示例）
        dangerous_domains = {'example.com', 'malicious.com'}
        if parsed_url.netloc in dangerous_domains:
            return False
            
        return True

    def _get_response(self, url: str) -> requests.Response:
        """获取URL响应"""
        try:
            # 通用请求头
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Cache-Control': 'max-age=0'
            }
            
            response = self.session.get(
                url,
                headers=headers,
                timeout=30,
                allow_redirects=True,
                verify=True
            )
            
            # 检查状态码
            if response.status_code == 403:
                raise ValueError(f"访问被拒绝 (HTTP 403)，该网页可能需要登录或不允许访问")
            elif response.status_code == 404:
                raise ValueError(f"页面未找到 (HTTP 404)")
            elif response.status_code == 429:
                raise ValueError(f"请求过于频繁 (HTTP 429)，请稍后再试")
            elif response.status_code >= 400:
                raise ValueError(f"请求失败，HTTP状态码: {response.status_code}")
            
            response.raise_for_status()
            
            # 检查内容类型
            content_type = response.headers.get('content-type', '').lower()
            if not any(t in content_type for t in ['text/html', 'application/json', 'text/plain']):
                raise ValueError(f"不支持的内容类型: {content_type}")
            
            return response
            
        except requests.Timeout:
            raise ValueError("请求超时，请检查网络连接或稍后重试")
        except requests.ConnectionError:
            raise ValueError("网络连接错误，请检查网络连接")
        except requests.TooManyRedirects:
            raise ValueError("重定向次数过多，可能是无效的URL")
        except requests.HTTPError as e:
            raise ValueError(f"HTTP请求错误: {str(e)}")
        except Exception as e:
            raise ValueError(f"请求失败: {str(e)}")

    def _parse_web_content(self, response: requests.Response) -> str:
        """解析网页内容"""
        # 检测编码
        response.encoding = response.apparent_encoding
        
        # 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 移除不需要的元素
        for element in soup(["script", "style", "meta", "link", "noscript", "header", "footer", "nav"]):
            element.decompose()
        
        # 获取主要内容
        main_content = (
            soup.find('main') or 
            soup.find('article') or 
            soup.find('div', class_='content') or 
            soup
        )
        
        # 获取文本内容
        text = main_content.get_text(separator='\n', strip=True)
        
        # 清理文本
        lines = []
        for line in text.splitlines():
            line = line.strip()
            if line and len(line) > 1:  # 忽略单字符行
                lines.append(line)
        
        return '\n'.join(lines)

    def _filter_content(self, text: str) -> str:
        """过滤和清理内容"""
        # 敏感内容过滤
        text = self._filter_sensitive_content(text)
        
        # 广告过滤
        text = self._filter_advertisements(text)
        
        # 去除重复内容
        text = self._remove_duplicates(text)
        
        # 清理 HTML
        text = bleach.clean(text)
        
        # 转义特殊字符
        text = html.escape(text)
        
        return text

    def _filter_sensitive_content(self, text: str) -> str:
        """过滤敏感内容"""
        for word in self.SENSITIVE_WORDS:
            text = re.sub(r'\b' + re.escape(word) + r'\b', '[FILTERED]', text, flags=re.IGNORECASE)
        return text
    
    def _filter_advertisements(self, text: str) -> str:
        """过滤广告"""
        for pattern in self.AD_PATTERNS:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        return text
    
    def _remove_duplicates(self, text: str) -> str:
        """去除重复内容"""
        lines = text.split('\n')
        unique_lines = []
        seen = set()
        
        for line in lines:
            line_hash = hashlib.md5(line.encode()).hexdigest()
            if line_hash not in seen and line.strip():
                seen.add(line_hash)
                unique_lines.append(line)
        
        return '\n'.join(unique_lines)

    def _handle_error(self, error: Exception):
        """错误处理"""
        if isinstance(error, Timeout):
            logger.error("Request timed out")
        elif isinstance(error, RequestException):
            logger.error(f"Request failed: {str(error)}")
        elif isinstance(error, ValueError):
            logger.error(f"Invalid input: {str(error)}")
        else:
            logger.error(f"Unexpected error: {str(error)}", exc_info=True) 