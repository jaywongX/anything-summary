import os
from typing import List, Optional
from PyPDF2 import PdfReader
from docx import Document
# import pytesseract
# from PIL import Image
# import whisper
# import ffmpeg
# from modelscope.pipelines import pipeline
# from modelscope.utils.constant import Tasks

class FileProcessor:
    def __init__(self, upload_dir: str):
        self.upload_dir = upload_dir
        # 暂时注释掉音频处理初始化
        # self.asr_pipeline = pipeline(
        #     task=Tasks.auto_speech_recognition,
        #     model='damo/speech_paraformer-large_asr_nat-zh-cn-16k-common-vocab8404-pytorch'
        # )
    
    async def process_pdf(self, file_path: str) -> str:
        """处理PDF文件"""
        try:
            with open(file_path, 'rb') as file:
                pdf = PdfReader(file)
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text.strip():  # 只添加非空内容
                        text_parts.append(text.strip())
                return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"PDF处理失败: {str(e)}")
    
    async def process_docx(self, file_path: str) -> str:
        """处理Word文档"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 提取标题
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    text_parts.append(f"# {paragraph.text}")
                elif paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Word文档处理失败: {str(e)}")
    
    async def process_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
            
    async def process_image(self, file_path: str) -> str:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)
    
    async def process_audio(self, file_path: str) -> str:
        """处理音频文件（临时实现）"""
        return "音频处理功能暂未启用"
    
    async def process_video(self, file_path: str) -> str:
        """处理视频文件（临时实现）"""
        return "视频处理功能暂未启用" 